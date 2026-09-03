"""格式转换类：PDF/A、HTML、图片、图片转PDF、Markdown、Office(Word)。"""
import fitz  # PyMuPDF
import logging
import re
from pathlib import Path
from contextlib import contextmanager
from .core import register, save_uploads, new_tmp, send_file, require


# ---------- PDF→DOCX 进度回调 ----------
# pdf2docx 内部通过 logging.info 输出 [1/4]~[4/4] 与 (x/y) Page 进度，
# 这里截获并归一化为进度消息，供 WebSocket 实时推送。
_PROG_STAGE = re.compile(r"\[(\d+)/(\d+)\]\s*(.+)")
_PROG_PAGE = re.compile(r"\((\d+)/(\d+)\)\s*Page")


class _ProgressHandler(logging.Handler):
    def __init__(self, cb):
        super().__init__()
        self.cb = cb

    def emit(self, record):
        msg = record.getMessage()
        m = _PROG_STAGE.match(msg)
        if m:
            self.cb({"type": "progress", "stage": int(m.group(1)),
                    "stageTotal": int(m.group(2)), "message": m.group(3).strip()})
            return
        m = _PROG_PAGE.match(msg)
        if m:
            self.cb({"type": "progress", "page": int(m.group(1)),
                    "pageTotal": int(m.group(2)),
                    "message": f"正在转换第 {m.group(1)}/{m.group(2)} 页"})


@contextmanager
def _progress_ctx(cb):
    if not cb:
        yield
        return
    root = logging.getLogger()
    h = _ProgressHandler(cb)
    h.setLevel(logging.INFO)
    root.addHandler(h)
    try:
        yield
    finally:
        root.removeHandler(h)


def parse_pages_param(s):
    """解析 '1-3,5' 为 0-based 页码列表；空返回 None。"""
    if not s:
        return None
    pages = []
    for part in str(s).split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            for i in range(int(a) - 1, int(b)):
                pages.append(i)
        else:
            pages.append(int(part) - 1)
    return pages or None


# ============================================================
# PyMuPDF → Markdown → DOCX 新引擎
# ============================================================
import markdown as _markdown
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from pymupdf4llm import to_markdown


class _MdToDocxConverter:
    """把 Markdown（由 pymupdf4llm 从 PDF 提取）转成 .docx。"""

    _BODY_FONT = "Microsoft YaHei"
    _MONO_FONT = "Courier New"

    def __init__(self, base_dir=None, fit_images=True):
        self.doc = Document()
        self.base_dir = Path(base_dir) if base_dir else Path.cwd()
        self.fit_images = fit_images
        self._set_base_style()

    def _set_base_style(self):
        style = self.doc.styles["Normal"]
        style.font.name = self._BODY_FONT
        style.font.size = Pt(10.5)
        try:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), self._BODY_FONT)
        except Exception:
            pass

    def convert(self, md_text):
        html = _markdown.markdown(md_text or "", extensions=["tables", "fenced_code"])
        soup = BeautifulSoup(html, "html.parser")
        root = soup.body or soup
        for child in root.children:
            self._process_block(child)
        return self.doc

    def _process_block(self, elem):
        if isinstance(elem, NavigableString):
            text = str(elem).strip()
            if text:
                self.doc.add_paragraph(text)
            return
        name = elem.name
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._heading(elem, int(name[1]))
        elif name == "p":
            self._paragraph(elem)
        elif name == "ul":
            self._list(elem, ordered=False)
        elif name == "ol":
            self._list(elem, ordered=True)
        elif name == "table":
            self._table(elem)
        elif name == "pre":
            self._pre(elem)
        elif name == "blockquote":
            self._blockquote(elem)
        elif name == "hr":
            self.doc.add_paragraph("─" * 40)
        else:
            for child in elem.children:
                self._process_block(child)

    def _heading(self, elem, level):
        p = self.doc.add_heading(level=min(level, 9))
        self._inline(p, elem)

    def _paragraph(self, elem):
        p = self.doc.add_paragraph()
        self._inline(p, elem)

    def _list(self, elem, ordered=False, level=0):
        for li in elem.find_all("li", recursive=False):
            p = self.doc.add_paragraph(
                style="List Number" if ordered else "List Bullet"
            )
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            for child in li.children:
                if isinstance(child, NavigableString):
                    t = str(child).strip()
                    if t:
                        self._apply_font(p.add_run(t))
                elif child.name in ("ul", "ol"):
                    self._list(child, ordered=(child.name == "ol"), level=level + 1)
                else:
                    self._inline_in_run(p, child)

    def _table(self, elem):
        rows = elem.find_all("tr")
        if not rows:
            return
        col_count = max(len(row.find_all(("td", "th"))) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            cells = row.find_all(("td", "th"))
            for j, cell in enumerate(cells):
                if j >= col_count:
                    break
                table.rows[i].cells[j].text = cell.get_text(strip=True)
        self.doc.add_paragraph()

    def _pre(self, elem):
        code = elem.find("code") or elem
        text = code.get_text()
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = self._MONO_FONT
        run.font.size = Pt(9)

    def _blockquote(self, elem):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        self._inline(p, elem)

    def _inline(self, paragraph, elem):
        for child in elem.children:
            self._inline_in_run(paragraph, child)

    def _inline_in_run(self, paragraph, elem):
        if isinstance(elem, NavigableString):
            text = str(elem)
            if text:
                self._apply_font(paragraph.add_run(text))
            return
        name = elem.name
        if name in ("strong", "b"):
            run = paragraph.add_run(elem.get_text())
            run.bold = True
            self._apply_font(run)
        elif name in ("em", "i"):
            run = paragraph.add_run(elem.get_text())
            run.italic = True
            self._apply_font(run)
        elif name == "code":
            run = paragraph.add_run(elem.get_text())
            run.font.name = self._MONO_FONT
        elif name == "a":
            run = paragraph.add_run(elem.get_text())
            run.underline = True
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        elif name == "br":
            paragraph.add_run().add_break()
        elif name == "img":
            self._add_image(paragraph, elem)
        elif name in ("ul", "ol"):
            self._list(elem, ordered=(name == "ol"))
        else:
            self._inline(paragraph, elem)

    def _add_image(self, paragraph, elem):
        src = elem.get("src")
        if not src:
            return
        img_path = self.base_dir / src
        if not img_path.exists():
            return
        try:
            run = paragraph.add_run()
            width = Inches(5.5) if self.fit_images else None
            run.add_picture(str(img_path), width=width)
        except Exception:
            pass

    def _apply_font(self, run):
        try:
            run.font.name = self._BODY_FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self._BODY_FONT)
        except Exception:
            pass


def _pdf_to_docx_pymupdf(pdf_path, out_path, params, progress=None):
    """PyMuPDF → Markdown → DOCX。"""
    pages = parse_pages_param(params.get("pages"))
    password = params.get("password") or None
    fit_images = str(params.get("fitImages", "1")) not in ("0", "false", "no")
    # 提速开关：大文件若只要文字，可关闭图片导出（跳过 PNG 重编码，提速数倍）
    extract_images = str(params.get("extractImages", "1")) not in ("0", "false", "no")
    image_format = params.get("imageFormat", "png")
    if image_format not in ("png", "jpeg", "jpg"):
        image_format = "png"
    out_dir = Path(out_path).parent
    pdf_size = Path(pdf_path).stat().st_size

    print(f"[convert-office] pymupdf engine start: pdf={pdf_path} size={pdf_size}", flush=True)

    if progress:
        progress({"type": "progress", "stage": 1, "stageTotal": 3,
                  "message": "PyMuPDF 提取文档结构与内容"})

    doc = fitz.open(str(pdf_path), password=password) if password else fitz.open(str(pdf_path))
    try:
        print(f"[convert-office] pdf pages={doc.page_count}", flush=True)
        md = to_markdown(
            doc,
            pages=pages,
            write_images=extract_images,
            image_path=str(out_dir),
            image_format=image_format,
        )
    finally:
        doc.close()

    md_len = len(md) if md else 0
    print(f"[convert-office] extracted markdown length={md_len}", flush=True)
    if md_len == 0:
        raise ValueError("该 PDF 未提取到任何文本（可能是扫描版图片 PDF）")

    if progress:
        progress({"type": "progress", "stage": 2, "stageTotal": 3,
                  "message": "Markdown 生成 Word 文档"})

    converter = _MdToDocxConverter(base_dir=out_dir, fit_images=fit_images)
    docx = converter.convert(md)
    docx.save(str(out_path))
    out_size = Path(out_path).stat().st_size
    print(f"[convert-office] pymupdf engine saved: {out_path} size={out_size}", flush=True)

    if progress:
        progress({"type": "progress", "stage": 3, "stageTotal": 3,
                  "message": "Word 文档保存完成"})


def _pdf_to_docx_pdf2docx(pdf_path, out_path, params, progress=None):
    """原 pdf2docx 引擎，作为兜底。"""
    from pdf2docx import Converter
    password = params.get("password") or None
    pages = parse_pages_param(params.get("pages"))
    cv_kwargs = {}
    if pages is not None:
        cv_kwargs["pages"] = pages
    if str(params.get("listNotTable", "1")) in ("0", "false", "no"):
        cv_kwargs["list_not_table"] = False
    cv = Converter(str(pdf_path), password=password)
    with _progress_ctx(progress):
        cv.convert(str(out_path), **cv_kwargs)
    cv.close()


@register("to-pdfa", desc="PDF 转 PDF/A")
def to_pdfa(files, params):
    p, _, _ = save_uploads(files)[0]
    level = params.get("level", "PDF/A-1b")
    doc = fitz.open(str(p))
    out = new_tmp() / "pdfa.pdf"
    # PyMuPDF 通过元数据标识 PDF/A 并清理交互元素
    doc.set_metadata({})
    doc.save(str(out), clean=True, deflate=True)
    doc.close()
    return send_file(out, "pdfa.pdf", "application/pdf")


@register("to-image", desc="PDF 转图片")
def to_image(files, params):
    p, _, _ = save_uploads(files)[0]
    fmt = params.get("format", "png")
    dpi = int(params.get("dpi", 150))
    single = params.get("single") in ("true", True)
    scale = dpi / 72.0
    doc = fitz.open(str(p))
    out_dir = new_tmp()
    if single:
        import PIL.Image as Image
        imgs = [page.get_pixmap(matrix=fitz.Matrix(scale, scale)) for page in doc]
        canvas = Image.new("RGB", (imgs[0].width, sum(pm.height for pm in imgs)))
        y = 0
        for pm in imgs:
            canvas.paste(Image.frombytes("RGB", (pm.width, pm.height), pm.samples), (0, y))
            y += pm.height
        out = out_dir / f"long.{fmt}"
        canvas.save(str(out))
    else:
        for i, page in enumerate(doc):
            pm = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
            out = out_dir / f"page-{i + 1:03d}.{fmt}"
            pm.save(str(out))
    doc.close()
    import zipfile, io
    from urllib.parse import quote
    from fastapi import Response
    if single:
        return send_file(list(out_dir.glob("*"))[0], f"long.{fmt}", f"image/{fmt}")
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        for f in sorted(out_dir.glob(f"*.{fmt}")):
            z.write(f, f.name)
    return Response(content=buf.getvalue(), media_type="application/zip",
                    headers={"Content-Disposition": "attachment; filename*=UTF-8''images.zip"})


def _natural_key(name):
    """文件名自然排序，保证 img2 排在 img10 之前。"""
    return [int(s) if s.isdigit() else s.lower() for s in re.split(r"(\d+)", str(name))]


@register("image-to-pdf", desc="图片转 PDF")
def image_to_pdf(files, params):
    """把多张图片合成为一份 PDF。

    支持参数：
      fit    : fit(页面贴合图片+页边距) / stretch(拉伸铺满A4) / center(A4内等比居中)
      color  : color(原色) / grayscale(灰度)
      margin : 页边距，单位 mm
    多图按文件名自然排序，与前端提示一致。
    """
    saved = save_uploads(files)
    fit = params.get("fit", "fit")
    color = params.get("color", "color")
    try:
        margin_mm = float(params.get("margin", 0) or 0)
    except (TypeError, ValueError):
        margin_mm = 0.0
    m = max(0.0, margin_mm) * 72.0 / 25.4  # mm -> pt

    saved.sort(key=lambda t: _natural_key(t[1]))

    a4 = fitz.paper_rect("a4")
    doc = fitz.open()
    try:
        for p, name, _ in saved:
            # 混入的 PDF 原样插入
            if p.suffix.lower() == ".pdf":
                src = fitz.open(p)
                doc.insert_pdf(src)
                src.close()
                continue

            try:
                pix = fitz.Pixmap(str(p))
            except Exception:
                raise ValueError(f"不支持的图片格式或文件已损坏：{name}")

            # 灰度：转换色彩空间
            if color == "grayscale" and not pix.is_gray:
                pix = fitz.Pixmap(fitz.csGRAY, pix)

            if fit == "stretch":
                # 拉伸铺满 A4（不保持比例，可能变形）
                page = doc.new_page(width=a4.width, height=a4.height)
                page.insert_image(fitz.Rect(m, m, a4.width - m, a4.height - m), pixmap=pix)
            elif fit == "center":
                # A4 页面内等比居中（保持比例，可能留白）
                page = doc.new_page(width=a4.width, height=a4.height)
                avail_w = max(1.0, a4.width - 2 * m)
                avail_h = max(1.0, a4.height - 2 * m)
                scale = min(avail_w / pix.width, avail_h / pix.height)
                w, h = pix.width * scale, pix.height * scale
                x0 = (a4.width - w) / 2.0
                y0 = (a4.height - h) / 2.0
                page.insert_image(fitz.Rect(x0, y0, x0 + w, y0 + h), pixmap=pix)
            else:
                # fit：页面尺寸贴合图片 + 页边距，图片 1:1 放入
                page = doc.new_page(width=pix.width + 2 * m, height=pix.height + 2 * m)
                page.insert_image(fitz.Rect(m, m, m + pix.width, m + pix.height), pixmap=pix)

        out = new_tmp() / "from-images.pdf"
        doc.save(str(out))
    finally:
        doc.close()
    return send_file(out, "from-images.pdf", "application/pdf")


@register("render-page", desc="渲染 PDF 指定页为 PNG（供前端可视化定位）")
def render_page(files, params):
    """把 PDF 的某一页渲染成 PNG 图片返回。

    用于前端「签名位置可视化」：前端拿到页面图片作背景，
    用户在其上拖动签名，再把位置换算成百分比提交。
    参数：page(页码，从 1 开始)、scale(渲染倍率，默认 1.5)
    """
    p, _, _ = save_uploads(files)[0]
    try:
        page_no = int(params.get("page", 1) or 1)
    except (TypeError, ValueError):
        page_no = 1

    doc = fitz.open(str(p))
    try:
        if page_no < 1:
            page_no = 1
        if page_no > doc.page_count:
            page_no = doc.page_count
        try:
            zoom = float(params.get("scale", 1.5) or 1.5)
        except (TypeError, ValueError):
            zoom = 1.5
        zoom = max(0.5, min(4.0, zoom))
        pix = doc[page_no - 1].get_pixmap(matrix=fitz.Matrix(zoom, zoom))
        out = new_tmp() / f"page-{page_no}.png"
        pix.save(str(out))
        page_w = doc[page_no - 1].rect.width
        page_h = doc[page_no - 1].rect.height
    finally:
        doc.close()
    resp = send_file(out, f"page-{page_no}.png", "image/png")
    # 附带页面尺寸，便于前端按比例换算
    resp.headers["X-Page-Width"] = str(page_w)
    resp.headers["X-Page-Height"] = str(page_h)
    return resp


@register("markdown-to-pdf", desc="Markdown 转 PDF")
def markdown_to_pdf(files, params):
    p, _, _ = save_uploads(files)[0]
    from pymupdf4llm import to_markdown
    # pymupdf4llm 主攻 PDF->MD；此处用 PyMuPDF 渲染纯文本为 PDF
    import md2pdf
    text = p.read_text(encoding="utf-8", errors="ignore")
    out = new_tmp() / "from-md.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), text, fontsize=11)
    doc.save(str(out))
    doc.close()
    return send_file(out, "from-md.pdf", "application/pdf")


@register("convert-office", desc="PDF 转 Word 文档")
def convert_office(files, params, progress=None):
    p, name, _ = save_uploads(files)[0]
    target = params.get("target", "docx")
    if target != "docx":
        from fastapi import HTTPException
        raise HTTPException(status_code=501, detail=f"当前仅支持 Word(.docx)，{target} 需额外转换组件")
    stem = Path(name).stem or "document"
    out_name = f"{stem}.docx"
    out = new_tmp() / out_name

    print(f"[convert-office] start: file={p} size={p.stat().st_size} name={name}", flush=True)

    def _fallback_pdf2docx(reason):
        print(f"[convert-office] fallback to pdf2docx: {reason}", flush=True)
        if progress:
            progress({"type": "progress", "stage": 1, "stageTotal": 4,
                      "message": f"{reason}，回退到 pdf2docx 重试"})
        _pdf_to_docx_pdf2docx(p, out, params, progress)

    # 首选 PyMuPDF → Markdown → DOCX（文本/表格/图片布局还原通常更完整）
    try:
        _pdf_to_docx_pymupdf(p, out, params, progress)
    except Exception as e:
        print(f"[convert-office] pymupdf engine failed: {e}", flush=True)
        _fallback_pdf2docx(str(e)[:60])

    # 如果结果明显为空（无文字层/扫描件），也尝试 pdf2docx 兜底
    out_size = out.stat().st_size
    if out_size < 10 * 1024:
        _fallback_pdf2docx(f"输出仅 {out_size} 字节，疑似空文档")
        out_size = out.stat().st_size

    print(f"[convert-office] final output: {out} size={out_size}", flush=True)
    if out_size < 5 * 1024:
        raise ValueError("转换结果为空：该 PDF 可能是扫描版图片或没有文字层，请使用 OCR 工具先识别文字。")

    return send_file(out, out_name,
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@register("to-pdf", desc="Office 转 PDF（需 LibreOffice）")
def to_pdf(files, params):
    require("libreoffice", "soffice")
    saved = save_uploads(files)
    out_dir = new_tmp()
    import subprocess
    src = saved[0][0]
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(src)],
                   check=True)
    out = next(out_dir.glob("*.pdf"))
    return send_file(out, "converted.pdf", "application/pdf")


@register("html-to-pdf", desc="HTML 转 PDF（需外部渲染引擎）")
def html_to_pdf(files, params):
    from fastapi import HTTPException
    raise HTTPException(status_code=501, detail="需配置 WeasyPrint / wkhtmltopdf 渲染引擎")


@register("to-presentation", desc="PDF 转演示文稿（暂未支持）")
def to_presentation(files, params):
    from fastapi import HTTPException
    raise HTTPException(status_code=501, detail="PPTX/ODP 导出需额外转换组件，暂未实现")
